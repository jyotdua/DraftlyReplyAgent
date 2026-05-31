from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
REPORTS_DIR = ROOT / "reports"
DIAGRAMS_DIR = REPORTS_DIR / "diagrams"
WORD_DIR = REPORTS_DIR / "word"
MARKDOWN_PATH = REPORTS_DIR / "draftly-report.md"
OUTPUT_DOCX = WORD_DIR / "draftly-report.docx"
HLD_IMAGE = DIAGRAMS_DIR / "draftly-hld.png"
LLD_IMAGE = DIAGRAMS_DIR / "draftly-lld-entities.png"

PAGE_WIDTH_IN = 8.27
PAGE_HEIGHT_IN = 11.69
MARGIN_IN = 0.65
USABLE_WIDTH_IN = PAGE_WIDTH_IN - (2 * MARGIN_IN)

BG = "#fcf7f2"
PRIMARY = "#ff5c00"
DARK = "#2a1b16"
BORDER = "#c7b6a7"
LIGHT = "#fffaf5"
CYLINDER = "#f4e8dc"
TEXT = "#231815"
MUTED = "#8b624e"


def ensure_dirs() -> None:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
                "/System/Library/Fonts/Supplemental/Helvetica Bold.ttf",
                "/Library/Fonts/Arial Bold.ttf",
            ]
        )
    candidates.extend(
        [
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/System/Library/Fonts/Supplemental/Helvetica.ttc",
            "/Library/Fonts/Arial.ttf",
        ]
    )
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], fill: str, outline: str, radius: int = 24, width: int = 3) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = DARK, width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    head_len = 18
    head_w = 10
    p1 = end
    p2 = (int(end[0] - head_len * ux + head_w * px), int(end[1] - head_len * uy + head_w * py))
    p3 = (int(end[0] - head_len * ux - head_w * px), int(end[1] - head_len * uy - head_w * py))
    draw.polygon([p1, p2, p3], fill=color)


def draw_multiline_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont, fill: str) -> None:
    left, top, right, bottom = box
    lines = text.split("\n")
    gaps = 8
    metrics = []
    total_height = 0
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        metrics.append((line, width, height))
        total_height += height
    total_height += gaps * (len(lines) - 1)
    cursor_y = top + ((bottom - top - total_height) // 2)
    for line, width, height in metrics:
        x = left + ((right - left - width) // 2)
        draw.text((x, cursor_y), line, font=font, fill=fill)
        cursor_y += height + gaps


def draw_multiline_left(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, font: ImageFont.ImageFont, fill: str, line_gap: int = 6) -> None:
    cursor_y = y
    for line in text.split("\n"):
        draw.text((x, cursor_y), line, font=font, fill=fill)
        bbox = draw.textbbox((x, cursor_y), line, font=font)
        cursor_y += (bbox[3] - bbox[1]) + line_gap


def generate_hld_png(output_path: Path) -> None:
    img = Image.new("RGB", (1500, 980), BG)
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, bold=True)
    node_font = get_font(26, bold=True)
    label_font = get_font(22)

    draw.text((60, 35), "Draftly High-Level Design", font=title_font, fill=DARK)
    draw.text((60, 85), "Backend-first Gmail AI reply architecture", font=label_font, fill=MUTED)

    boxes = {
        "client": (70, 400, 320, 500),
        "fastapi": (390, 390, 660, 510),
        "gmail": (790, 145, 1045, 245),
        "workflow": (760, 390, 1075, 515),
        "gemini": (790, 690, 1045, 790),
        "repo": (1170, 315, 1420, 415),
        "db": (1170, 590, 1420, 760),
    }

    rounded_box(draw, boxes["client"], PRIMARY, "#c94800", radius=28)
    draw_multiline_center(draw, boxes["client"], "Client / Postman /\nFuture UI", node_font, "#ffffff")

    rounded_box(draw, boxes["fastapi"], DARK, DARK, radius=28)
    draw_multiline_center(draw, boxes["fastapi"], "FastAPI Routes\n(app/main.py)", node_font, "#ffffff")

    for key in ("gmail", "workflow", "gemini", "repo"):
        rounded_box(draw, boxes[key], LIGHT, BORDER, radius=24)
    draw_multiline_center(draw, boxes["gmail"], "Gmail Service", node_font, TEXT)
    draw_multiline_center(draw, boxes["workflow"], "Draft Workflow Service", node_font, TEXT)
    draw_multiline_center(draw, boxes["gemini"], "Gemini Service", node_font, TEXT)
    draw_multiline_center(draw, boxes["repo"], "Repository Layer", node_font, TEXT)

    x1, y1, x2, y2 = boxes["db"]
    draw.rectangle((x1, y1 + 30, x2, y2 - 25), fill=CYLINDER, outline=BORDER, width=3)
    draw.ellipse((x1, y1, x2, y1 + 60), fill=CYLINDER, outline=BORDER, width=3)
    draw.ellipse((x1, y2 - 55, x2, y2), fill=CYLINDER, outline=BORDER, width=3)
    draw_multiline_center(draw, boxes["db"], "SQLite Database", node_font, TEXT)

    draw_multiline_center(
        draw,
        (740, 35, 1110, 105),
        "OAuth2, inbox fetch,\nthread fetch, draft create, send",
        label_font,
        MUTED,
    )
    draw_multiline_center(
        draw,
        (700, 560, 1125, 650),
        "Prompt composition, draft lifecycle,\nretry handling",
        label_font,
        MUTED,
    )

    arrow(draw, (320, 450), (390, 450))
    arrow(draw, (660, 430), (790, 195))
    arrow(draw, (660, 450), (760, 450))
    arrow(draw, (660, 470), (790, 740))
    arrow(draw, (1045, 450), (1170, 365))
    arrow(draw, (1045, 195), (1170, 365))
    arrow(draw, (920, 690), (920, 515))
    arrow(draw, (1295, 415), (1295, 590))

    img.save(output_path, "PNG")


def generate_lld_png(output_path: Path) -> None:
    img = Image.new("RGB", (1600, 1100), BG)
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, bold=True)
    entity_title_font = get_font(24, bold=True)
    entity_body_font = get_font(21)
    label_font = get_font(22, bold=True)

    draw.text((60, 35), "Draftly Low-Level Design: Entity Model", font=title_font, fill=DARK)
    draw.text((60, 85), "Core persistence objects and their relationships", font=entity_body_font, fill=MUTED)

    entities = {
        "user": (60, 150, 410, 395),
        "oauth": (500, 150, 840, 395),
        "draft": (50, 520, 500, 900),
        "send": (660, 560, 1030, 860),
        "audit": (1140, 600, 1490, 850),
    }

    for box in entities.values():
        rounded_box(draw, box, LIGHT, BORDER, radius=24)

    entity_text = {
        "user": ("UserAccount", ["PK: id", "email", "encrypted_credentials", "encrypted_preferences"]),
        "oauth": ("OAuthState", ["PK: id", "state", "redirect_uri", "FK: user_id", "user_email_hint"]),
        "draft": (
            "EmailDraft",
            [
                "PK: id",
                "FK: user_id",
                "source_message_id",
                "thread_id",
                "subject, tone, status",
                "generated_body, edited_body",
                "gmail_draft_id, gmail_sent_message_id",
                "prompt_context, draft_metadata",
            ],
        ),
        "send": (
            "SendAttempt",
            [
                "PK: id",
                "FK: draft_id",
                "attempt_number",
                "status",
                "idempotency_key",
                "gmail_message_id",
                "gmail_thread_id",
            ],
        ),
        "audit": ("AuditLog", ["PK: id", "FK: user_id", "event_type", "message", "payload"]),
    }

    for key, (title, fields) in entity_text.items():
        x1, y1, x2, y2 = entities[key]
        draw.text((x1 + 24, y1 + 18), title, font=entity_title_font, fill=DARK)
        draw.line((x1 + 22, y1 + 58, x2 - 22, y1 + 58), fill=BORDER, width=2)
        draw_multiline_left(draw, x1 + 24, y1 + 78, "\n".join(fields), entity_body_font, TEXT, line_gap=10)

    arrow(draw, (410, 270), (500, 270))
    arrow(draw, (235, 395), (235, 520))
    arrow(draw, (500, 710), (660, 710))
    draw.line((500, 850, 1080, 850), fill=DARK, width=4)
    arrow(draw, (1080, 850), (1140, 725))

    draw.text((430, 230), "1 : N", font=label_font, fill=MUTED)
    draw.text((255, 450), "1 : N", font=label_font, fill=MUTED)
    draw.text((565, 660), "1 : N", font=label_font, fill=MUTED)
    draw.text((690, 880), "user events", font=label_font, fill=MUTED)

    img.save(output_path, "PNG")


def set_page_margins(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(35, 24, 21)

    for style_name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11.5)):
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(42, 27, 22)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(10)
    r = p.add_run("Draftly Project Report")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(42, 27, 22)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(8)
    r = p.add_run("High-Level Design, Low-Level Design, Architecture, Workflow, and Security Overview")
    r.font.size = Pt(11.5)
    r.font.name = "Aptos"
    r.font.color.rgb = RGBColor(139, 98, 78)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(10)
    r = p.add_run("Prepared for the Draftly capstone submission")
    r.font.size = Pt(10.5)
    r.font.name = "Aptos"

    document.add_picture(str(HLD_IMAGE), width=Inches(5.9))
    last = document.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_after = Pt(10)

    add_image_caption(document, "System snapshot: client, FastAPI services, Gemini, Gmail, and database flow")

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Project Overview")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(42, 27, 22)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        "Draftly is a Gmail AI reply assistant built as a modular FastAPI backend. "
        "It connects to Gmail through OAuth2, reads email and thread context, generates "
        "tone-aware drafts with Gemini, stores workflow state in SQLite, and only sends "
        "replies after explicit user review and approval."
    )

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Report Highlights")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(42, 27, 22)

    for item in [
        "Problem statement, objectives, scope, and technology stack",
        "High-level architecture with service interaction flow",
        "Low-level design covering APIs, services, repositories, and entities",
        "End-to-end workflow for authentication, draft generation, review, and sending",
        "Security considerations, testing summary, and future enhancements",
    ]:
        add_bullet(document, item)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Core Stack: ")
    run.bold = True
    p.add_run("FastAPI, SQLAlchemy, SQLite, Gmail API, Google OAuth2, Gemini, Tenacity, Cryptography")

    document.add_page_break()


def add_image_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(139, 98, 78)


def parse_markdown_lines(markdown: str) -> list[str]:
    return markdown.splitlines()


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text.strip())


def add_numbered(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text.strip())


def add_codeblock(document: Document, lines: list[str]) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(lines):
        run = p.add_run(line if line else " ")
        run.font.name = "Menlo"
        run.font.size = Pt(9)
        if idx < len(lines) - 1:
            run.add_break()


def add_paragraph_with_bold(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    parts = re.split(r"(\*\*.*?\*\*)", text.strip())
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def build_docx() -> None:
    document = Document()
    set_page_margins(document)
    style_document(document)
    add_cover(document)

    footer = document.sections[0].footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)

    lines = parse_markdown_lines(MARKDOWN_PATH.read_text(encoding="utf-8"))
    in_codeblock = False
    code_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_codeblock:
                add_codeblock(document, code_lines)
                code_lines = []
                in_codeblock = False
            else:
                in_codeblock = True
            continue

        if in_codeblock:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            continue

        if stripped.startswith("## "):
            heading = stripped[3:]
            document.add_heading(heading, level=1)
            if heading == "8. High-Level Design (HLD)":
                document.add_picture(str(HLD_IMAGE), width=Inches(USABLE_WIDTH_IN))
                add_image_caption(document, "Figure 1. Draftly high-level architecture diagram")
            elif heading == "9. Low-Level Design (LLD)":
                document.add_picture(str(LLD_IMAGE), width=Inches(USABLE_WIDTH_IN))
                add_image_caption(document, "Figure 2. Draftly low-level entity design diagram")
            continue

        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
            continue

        if stripped.startswith("#### "):
            document.add_heading(stripped[5:], level=3)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_numbered(document, re.sub(r"^\d+\.\s+", "", stripped))
            continue

        if stripped.startswith("- "):
            add_bullet(document, stripped[2:])
            continue

        add_paragraph_with_bold(document, stripped)

    document.save(OUTPUT_DOCX)


def main() -> None:
    ensure_dirs()
    generate_hld_png(HLD_IMAGE)
    generate_lld_png(LLD_IMAGE)
    build_docx()
    print(f"Created {OUTPUT_DOCX}")
    print(f"Created {HLD_IMAGE}")
    print(f"Created {LLD_IMAGE}")


if __name__ == "__main__":
    main()
